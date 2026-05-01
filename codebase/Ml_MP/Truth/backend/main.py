import os
import sys
import time
from datetime import datetime
import re
from typing import List, Optional
import joblib
import pandas as pd
import numpy as np

from fastapi import FastAPI, HTTPException, File, UploadFile, Form
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from PIL import Image
import io
import PyPDF2
import wikipedia
import requests
import json
import torch

# Load .env (if present) so API keys live outside the repo. We look in both
# the backend dir and the project root, so a single .env can serve everyone.
try:
    from dotenv import load_dotenv
    _backend_env = os.path.join(os.path.dirname(os.path.abspath(__file__)), ".env")
    _root_env = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "..", "..", "..", ".env")
    for _candidate in (_backend_env, _root_env):
        if os.path.exists(_candidate):
            load_dotenv(_candidate)
            print(f"[ENV] Loaded {_candidate}")
            break
except Exception as _e:
    print(f"[ENV] python-dotenv unavailable; relying on shell env: {_e}")

GOOGLE_FACTCHECK_API_KEY = os.environ.get("GOOGLE_FACTCHECK_API_KEY", "").strip() or None

# Gemini API for the LLM Plausibility verifier. We deliberately use
# gemini-2.5-flash with thinking disabled — this is a label-extraction task
# (EXTRAORDINARY / PLAUSIBLE / UNCERTAIN), not a reasoning task, and Flash's
# default thinking budget would burn output tokens before the label appears.
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY", "").strip() or None
LLM_PLAUSIBILITY_MODEL = os.environ.get("LLM_PLAUSIBILITY_MODEL", "gemini-2.5-flash").strip()
_gemini_client = None

def _get_gemini():
    """Lazy-init the Gemini SDK client. Returns None if no key is configured."""
    global _gemini_client
    if _gemini_client is None and GEMINI_API_KEY:
        try:
            from google import genai
            _gemini_client = genai.Client(api_key=GEMINI_API_KEY)
        except Exception as e:
            print(f"[LLM] Gemini SDK unavailable: {e}")
            _gemini_client = False  # sentinel — distinguishes "not yet tried" from "failed"
    return _gemini_client or None
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse

# Persistent Cache for RAG (Avoid API bans & increase speed).
# Entries carry a `_ts` epoch-seconds timestamp so we can expire stale
# verifier results (fact-checks evolve, retracted articles get pulled, etc.).
CACHE_FILE = "rag_cache.json"
CACHE_TTL_SECONDS = 7 * 24 * 60 * 60  # 7 days; tune per category if needed

def load_cache():
    if os.path.exists(CACHE_FILE):
        try:
            with open(CACHE_FILE, "r") as f:
                return json.load(f)
        except Exception as e:
            print(f"[CACHE] Failed to read {CACHE_FILE}: {e}; starting empty.")
    return {}

RAG_CACHE = load_cache()

def save_cache():
    try:
        with open(CACHE_FILE, "w") as f:
            json.dump(RAG_CACHE, f)
    except Exception as e:
        print(f"[CACHE] Save failed: {e}")

def cache_get(category: str, key: str):
    """Return a cached value only if it exists AND is younger than CACHE_TTL_SECONDS.

    Stale entries are removed lazily so the cache file shrinks as well as grows.
    Backwards-compatible with the legacy schema where entries had no `_ts` —
    those are treated as fresh on first read and re-stamped on the next write.
    """
    bucket = RAG_CACHE.get(category)
    if not bucket or key not in bucket:
        return None
    entry = bucket[key]
    if not isinstance(entry, dict):
        return entry  # legacy raw value; assume fresh
    ts = entry.get("_ts")
    if ts is None:
        # Legacy entry — treat as fresh, but copy without the bookkeeping key
        return {k: v for k, v in entry.items() if k != "_ts"}
    if (time.time() - float(ts)) > CACHE_TTL_SECONDS:
        del bucket[key]
        return None
    return {k: v for k, v in entry.items() if k != "_ts"}

def cache_set(category: str, key: str, value: dict):
    """Store a value with a fresh timestamp and persist to disk."""
    bucket = RAG_CACHE.setdefault(category, {})
    bucket[key] = {**value, "_ts": time.time()}
    save_cache()

ocr_reader = None
def initialize_ocr():
    """Initializes the EasyOCR engine with GPU/CUDA detection."""
    global ocr_reader
    try:
        import easyocr
        use_gpu = torch.cuda.is_available()
        ocr_reader = easyocr.Reader(['en'], gpu=use_gpu)
        print(f"EasyOCR Engine Online. GPU Acceleration: {use_gpu}")
    except Exception as e:
        print(f"EasyOCR Init failed: {e}")
        ocr_reader = None

# Boot the engine in the background
import threading
threading.Thread(target=initialize_ocr, daemon=True).start()

try:
    from duckduckgo_search import DDGS
except ImportError:
    pass

from urllib.parse import urlparse

# ==========================================
# SOURCE CREDIBILITY TIERS
# ==========================================
DOMAIN_TIERS = {
    "tier1": {
        "reuters.com", "apnews.com", "bbc.com", "bbc.co.uk",
        "nytimes.com", "washingtonpost.com", "theguardian.com",
        "nature.com", "science.org", "thelancet.com", "who.int",
        "un.org", "worldbank.org",
    },
    "tier2": {
        "cnn.com", "aljazeera.com", "dw.com", "france24.com",
        "ndtv.com", "thehindu.com", "hindustantimes.com",
        "timesofindia.indiatimes.com", "indianexpress.com",
        "economictimes.indiatimes.com", "livemint.com",
        "abc.net.au", "npr.org", "pbs.org", "usatoday.com",
        "forbes.com", "bloomberg.com",
    },
    "factcheck": {
        "snopes.com", "politifact.com", "factcheck.org", "fullfact.org",
        "altnews.in", "boomlive.in", "vishvasnews.com",
    },
}

def _get_domain(url: str) -> str:
    """Extract domain from URL."""
    try:
        domain = urlparse(url).netloc.lower()
        if domain.startswith("www."):
            domain = domain[4:]
        return domain
    except Exception:
        return ""

def _get_domain_credibility(url: str) -> tuple:
    """Returns (tier_name, credibility_score) for a URL."""
    domain = _get_domain(url)
    if domain.endswith(".gov") or domain.endswith(".edu"):
        return ("tier1", 0.95)
    for tier_name, domains in DOMAIN_TIERS.items():
        for d in domains:
            if d in domain:
                scores = {"tier1": 0.90, "tier2": 0.75, "factcheck": 0.95}
                return (tier_name, scores[tier_name])
    return ("unranked", 0.40)

_STOPWORDS = {"the", "a", "an", "is", "are", "was", "were", "be", "been",
              "being", "have", "has", "had", "do", "does", "did", "will",
              "would", "could", "should", "may", "might", "shall", "can",
              "in", "on", "at", "to", "for", "of", "with", "by", "from",
              "and", "or", "but", "not", "no", "if", "then", "than",
              "that", "this", "it", "its", "as", "so", "up", "about",
              "he", "she", "they", "we", "i", "you", "my", "his", "her"}

def _ddgs_search(query: str, max_results: int = 8) -> list:
    """DuckDuckGo text search with retry on rate limit."""
    for attempt in range(3):
        try:
            with DDGS() as ddgs:
                return list(ddgs.text(query, max_results=max_results))
        except Exception as e:
            if "Ratelimit" in str(e) and attempt < 2:
                time.sleep(2 * (attempt + 1))
                continue
            raise
    return []

def _extract_search_queries(text: str) -> list:
    """Generate multiple search queries from text for broader coverage."""
    sentences = [s.strip() for s in re.split(r'[.!?]\s+', text) if len(s.strip()) > 10]
    queries = []
    if sentences:
        queries.append(sentences[0][:150])
    try:
        from src.preprocessor import _load_spacy
        nlp = _load_spacy()
        doc = nlp(text[:500])
        entities = [ent.text for ent in doc.ents if ent.label_ in ("PERSON", "ORG", "GPE", "EVENT")]
        if entities:
            entity_query = " ".join(entities[:3])
            verbs = [tok.text for tok in doc if tok.pos_ == "VERB" and len(tok.text) > 3]
            if verbs:
                entity_query += " " + verbs[0]
            queries.append(entity_query[:150])
    except Exception:
        pass
    if len(sentences) > 1:
        queries.append(sentences[1][:150])
    return queries[:3]

# ==========================================
# LOAD TRUTHLENS AI CORES ON BOOT
# ==========================================
backend_dir = os.path.dirname(os.path.abspath(__file__))
truth_dir = os.path.dirname(backend_dir)
root_dir = os.path.dirname(truth_dir)
truthlens_dir = os.path.join(root_dir, "TruthLens")

# Inject TruthLens into Python Path to access the custom FeatureEngine class
sys.path.append(truthlens_dir)

ensemble_model = None
feature_engine = None

try:
    from src.feature_engineer import TruthLensFeatureEngine
    from src.preprocessor import preprocess_dataframe
    from src.claim_detector import classify_input as classify_claim_input
    import config as truthlens_config
    try:
        from scripts.sentence_scorer import score_sentences  # type: ignore
    except Exception as _ss_e:
        score_sentences = None
        print(f"[BOOT] sentence_scorer unavailable: {_ss_e}")
    TRUTHLENS_MODELS_DIR = os.path.join(truthlens_dir, "models")
    print(f"Loading TruthLens AI Cores from: {TRUTHLENS_MODELS_DIR}")

    ensemble_model = joblib.load(os.path.join(TRUTHLENS_MODELS_DIR, "ensemble.joblib"))
    feature_engine = TruthLensFeatureEngine().load(os.path.join(TRUTHLENS_MODELS_DIR, "feature_engine.pkl"))

    print("TruthLens AI Meta-Learner Successfully Loaded into Server Memory!")
except Exception as e:
    truthlens_config = None
    classify_claim_input = None
    score_sentences = None
    print(f"CRITICAL: Failed to load TruthLens Cores. Using Text Heuristics. Error: {e}")

# Threshold for the vocabulary-coverage abstention gate. Pulled from
# TruthLens config when available; falls back to the historical default.
_VOCAB_COVERAGE_THRESHOLD = (
    getattr(truthlens_config, "VOCAB_COVERAGE_THRESHOLD", 0.30)
    if truthlens_config else 0.30
)
_MIN_MATCHED_TERMS = (
    getattr(truthlens_config, "MIN_MATCHED_TERMS", 4)
    if truthlens_config else 4
)

# ==========================================

app = FastAPI(
    title="Truth API",
    description="Backend API for the Truth fake news detection platform with internal ML inference and Web RAG.",
    version="1.0.0"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://localhost:3000", "http://localhost:8000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

class LimeFeature(BaseModel):
    word: str
    weight: float
    color: str

class FactualCheck(BaseModel):
    check_type: str
    result: str
    status: str # "verified", "conflict", "unknown"
    url: Optional[str] = None

class SentenceScore(BaseModel):
    """Per-sentence verdict for Originality.ai-style highlighting."""
    sentence: str
    prob_fake: float
    prob_real: float
    risk: str   # "low" | "medium" | "high"
    char_start: int
    char_end: int

class SourceCredibility(BaseModel):
    """NewsGuard-style source-tier surfacing for any URL in the input."""
    domain: str
    tier: str    # "tier1" | "tier2" | "factcheck" | "unranked"
    score: float # 0.0–1.0

class ConflictReport(BaseModel):
    """Structured ML-vs-RAG agreement breakdown — drives the Conflict
    Explainer panel in the Analyzer UI. Populated by build_conflict_report().
    """
    ml_verdict: str
    ml_confidence: float
    rag_verdict: str
    rag_confidence: float
    disagreement: bool
    winning_signal: str       # "ml" | "rag" | "consensus"
    triggered_rule: Optional[str] = None
    flagging_verifiers: List[str] = []
    top_lime_tokens: List[LimeFeature] = []
    vocab_coverage: float = 1.0
    bias_gate_active: bool = False

class AnalyzeResponse(BaseModel):
    prediction: str
    confidence: float
    processing_time: float
    explainability: List[LimeFeature]
    analyzed_text: str = ""
    factual_analysis: List[FactualCheck] = []
    conflict_detected: bool = False
    vocab_coverage: float = 1.0  # How much of the input the model understands (0-1)
    # Parity surfacing (Phase 1)
    sentence_scores: List[SentenceScore] = []
    claim_type: Optional[str] = None
    source_credibility: Optional[SourceCredibility] = None
    # Self-auditing surfacing (Phase 4)
    conflict_report: Optional[ConflictReport] = None

def verify_timeline(text: str) -> FactualCheck:
    """Verifies that dates are valid using system calendar and timezones."""
    from datetime import timedelta, timezone
    import calendar as cal_module
    
    now = datetime.now()
    current_year = now.year
    current_month = now.month
    current_day = now.day
    text_lower = text.lower()

    # Calendar integration: Check for valid days in months
    days_in_month = {
        1: 31, 2: 29 if (current_year % 4 == 0 and (current_year % 100 != 0 or current_year % 400 == 0)) else 28,
        3: 31, 4: 30, 5: 31, 6: 30, 7: 31, 8: 31, 9: 30, 10: 31, 11: 30, 12: 31
    }
    
    yesterday = now - timedelta(days=1)
    tomorrow = now + timedelta(days=1)
    
    # Extract date patterns: "April 31" (invalid), "February 30" (invalid), etc.
    month_names = ['january', 'february', 'march', 'april', 'may', 'june', 
                   'july', 'august', 'september', 'october', 'november', 'december']
    
    for month_idx, month_name in enumerate(month_names, 1):
        # Use negative lookahead to avoid matching year digits as day numbers
        # e.g. "april 2026" should NOT capture "20" as a day
        pattern = rf'{month_name}\s+(\d{{1,2}})(?!\d)'
        matches = re.finditer(pattern, text_lower)
        for match in matches:
            day = int(match.group(1))
            max_day = days_in_month[month_idx]
            if day > max_day:
                return FactualCheck(
                    check_type="Temporal Validation",
                    result=f"Conflict: {month_name.capitalize()} only has {max_day} days, but claim references day {day}.",
                    status="conflict"
                )

    # Timeline consistency check: "Today is [date]" — validates day, month, AND year
    if "today" in text_lower:
        # Match patterns like "is today 7th april 2026" or "today is april 7, 2026"
        match_dmy = re.search(r'(?:today\s+is|is\s+today)\s+(\d{1,2})(?:st|nd|rd|th)?\s+(\w+)\s+(\d{4})', text_lower)
        match_mdy = re.search(r'(?:today\s+is|is\s+today)\s+(\w+)\s+(\d{1,2})(?:st|nd|rd|th)?\s*,?\s*(\d{4})', text_lower)

        day_m = month_m = year_m = None
        if match_dmy:
            day_m = int(match_dmy.group(1))
            month_str = match_dmy.group(2)
            year_m = int(match_dmy.group(3))
            month_m = (month_names.index(month_str) + 1) if month_str in month_names else None
        elif match_mdy:
            month_str = match_mdy.group(1)
            day_m = int(match_mdy.group(2))
            year_m = int(match_mdy.group(3))
            month_m = (month_names.index(month_str) + 1) if month_str in month_names else None

        if day_m and month_m and year_m:
            if year_m != current_year or month_m != current_month or day_m != current_day:
                claimed = f"{month_names[month_m-1].capitalize()} {day_m}, {year_m}"
                actual = f"{month_names[current_month-1].capitalize()} {current_day}, {current_year}"
                return FactualCheck(
                    check_type="Temporal Validation",
                    result=f"Conflict: Claims today is {claimed}, but actual date is {actual}.",
                    status="conflict"
                )
            else:
                return FactualCheck(
                    check_type="Temporal Validation",
                    result=f"Verified: Today's date confirmed as {month_names[current_month-1].capitalize()} {current_day}, {current_year}.",
                    status="verified"
                )

    # Check for future years
    years_mentioned = re.findall(r'\b(20\d{2})\b', text)
    if years_mentioned:
        for year in years_mentioned:
            if int(year) > current_year + 10:  # Allow up to 10 years in future (reasonable for planning)
                return FactualCheck(
                    check_type="Temporal Validation",
                    result=f"Conflict: Claim mentions year {year}, which is too far in the future.",
                    status="conflict"
                )
        return FactualCheck(
            check_type="Temporal Validation",
            result=f"Valid: Calendar checks passed for year {years_mentioned[0]}.",
            status="verified"
        )
    return FactualCheck(check_type="Temporal Validation", result="No specific dates found to verify.", status="unknown")

def verify_wikipedia(text: str) -> FactualCheck:
    """Uses Wikipedia API to find relevant summaries and check for major factual conflicts."""
    # 0. Check Cache (TTL-respecting)
    import hashlib
    cache_key = hashlib.md5(text.strip().encode()).hexdigest()
    cached = cache_get("wikipedia", cache_key)
    if cached is not None:
        return FactualCheck(**cached)

    try:
        # 1. Subject Extraction using spaCy NER (not regex)
        # spaCy NER properly identifies PERSON, ORG, GPE, LOC entities
        try:
            from src.preprocessor import _load_spacy
            nlp = _load_spacy()
            doc = nlp(text)
            ner_subjects = []
            for ent in doc.ents:
                if ent.label_ in ("PERSON", "ORG", "GPE", "LOC", "FAC", "NORP", "EVENT", "WORK_OF_ART"):
                    if len(ent.text.strip()) > 2:
                        ner_subjects.append(ent.text.strip())
            # Deduplicate, longest first
            valid_subjects = sorted(list(set(ner_subjects)), key=len, reverse=True)
        except Exception:
            # Fallback to regex if spaCy unavailable in this context
            found_subjects = re.findall(r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)+\b', text)
            stop_titles = {"The", "A", "An", "On", "In", "Today", "Yesterday", "Tomorrow"}
            valid_subjects = [s for s in found_subjects if s not in stop_titles and len(s) > 4]

        # Candidate subjects for the Wikipedia link
        candidates = valid_subjects[:5] if valid_subjects else []
        if not candidates:
            return FactualCheck(check_type="Wikipedia Link", result="No matching encyclopedic record found.", status="unknown")
        
        all_checks = []
        for qc in candidates:
            try:
                # Use sub-query for better matching
                search_results = wikipedia.search(qc)
                if not search_results: continue
                
                found_query = search_results[0]
                summary = wikipedia.summary(found_query, sentences=3)
                sum_low = summary.lower()
                low_text = text.lower()
                
                # Check for conspiracy/hoax signatures OR Extraordinary Claims
                negative_sigs = ["conspiracy", "hoax", "debunked", "false claim", "pseudoscience", "misinformation"]
                # 'Myth' is only a conflict if it's NOT a physical monument (temple, statue, church)
                is_monument = any(kw in sum_low for kw in ["temple", "statue", "monument", "historic", "church", "mosque", "shrine", "cathedral", "unesco"])
                has_modern_date = any(yr in sum_low for yr in ["2024", "2023", "2022", "2018", "2011", "1919"])
                
                hallucination_sigs = ["teleportation", "flat earth", "antigravity", "perpetual motion", "water-powered engine", "free energy"]

                if any(ns in sum_low for ns in negative_sigs) or \
                   (any(hs in low_text for hs in hallucination_sigs) and ("theoretical" in sum_low or "process" in sum_low or "concept" in sum_low or "technique" in sum_low or "physics" in sum_low or "theory" in sum_low)):
                    all_checks.append(FactualCheck(
                        check_type="Wikipedia Link", 
                        result=f"Conflict: Subject '{found_query}' is flagged as a myth or misinformation.", 
                        status="conflict"
                    ))
                    continue

                # Rule: If 'myth' is present but it's a monument or has a modern date, it's NOT a conflict.
                if "myth" in sum_low and not is_monument and not has_modern_date:
                     all_checks.append(FactualCheck(
                        check_type="Wikipedia Link", 
                        result=f"Conflict: Subject '{found_query}' is flagged as mythological/non-physical.", 
                        status="conflict"
                    ))
                     continue

                # Context Lock
                cities_in_text = re.findall(r'\b(Mumbai|Delhi|New York|London|Paris|Borivali|Pune|Bangalore)\b', text, re.IGNORECASE)
                if cities_in_text and not any(c.lower() in sum_low for c in cities_in_text):
                     if any(hs in low_text for hs in hallucination_sigs):
                        all_checks.append(FactualCheck(
                            check_type="Wikipedia Link", 
                            result=f"Conflict: '{found_query}' has no verified record in {cities_in_text[0]}.", 
                            status="conflict"
                        ))
                        continue

                # Factual Numeric Check for Wikipedia
                wiki_numbers = re.findall(r'\b\d+\b', summary)
                text_numbers = re.findall(r'\b\d+\b', text)
                if text_numbers and wiki_numbers:
                    t_val = int(text_numbers[0])
                    if t_val == 0 and any(int(wn) > 0 for wn in wiki_numbers[:5]):
                         all_checks.append(FactualCheck(
                            check_type="Wikipedia Link", 
                            result=f"Conflict: Wikipedia record for '{found_query}' contradicts the zero-count claim.", 
                            status="conflict"
                        ))
                         continue

                # Final verification match
                if found_query.lower() in text.lower():
                    all_checks.append(FactualCheck(
                        check_type="Wikipedia Link", 
                        result=f"Verified: Record match ({found_query}).", 
                        status="verified",
                        url=f"https://en.wikipedia.org/wiki/{found_query.replace(' ', '_')}"
                    ))
            except Exception:
                continue

        if not all_checks:
            return FactualCheck(check_type="Wikipedia Link", result="No matching encyclopedic record found.", status="unknown")
            
        # THE MASTER ARCHITECT RULE: IF ANY SUBJECT CONFLICTS, THE WHOLE CHECK CONFLICTS.
        if any(c.status == "conflict" for c in all_checks):
            conflicts = [c for c in all_checks if c.status == "conflict"]
            return conflicts[0]
            
        # Otherwise, if any verified, return the strongest verification
        verified = [c for c in all_checks if c.status == "verified"]
        if verified:
            return verified[0]
            
        return FactualCheck(check_type="Wikipedia Link", result="Unknown: Records found but context mismatch.", status="unknown")
            
    except Exception as e:
        print(f"[RAG] Wikipedia error: {e}")
        return FactualCheck(check_type="Wikipedia Link", result="Wikipedia Engine offline.", status="unknown")

def scrub_noise(text: str) -> str:
    """Removes common Google Maps or Newspaper UI elements and OCR garbage strings."""
    # 1. Remove URLs and long-string query parameters
    text = re.sub(r'http\S+', '', text)
    text = re.sub(r'www\S+', '', text)
    text = re.sub(r'google\.com/search\?\S+', '', text)
    
    # 2. Aggressive Garbage Removal (Like '31Iqri4isi') - Remove words that are gibberish
    # Rules: Random mix of letters/numbers without vowels, or too long without spaces
    words = text.split()
    clean_words = []
    for w in words:
        # If word is long and has no vowels and contains numbers, it's garbage
        has_vowel = any(v in w.lower() for v in 'aeiou')
        has_digit = any(d.isdigit() for d in w)
        if len(w) > 6 and not has_vowel and has_digit:
            continue
        clean_words.append(w)
    text = " ".join(clean_words)

    unwanted = [
        "See photos", "See outside", "Google reviews", "Website", "Directions", "Share", "Call", 
        "Write a review", "Edit", "Suggest an edit", "Business", "Maps", "Search Google", 
        "Image by", "By", "Photo by", "Save", "Add a label", "Send to your phone", "Google", "Search",
        "Reviews", "Overview", "Updates", "About", "Questions", "Answers"
    ]
    txt = text
    for u in unwanted:
        txt = re.sub(rf'\b{u}\b', '', txt, flags=re.IGNORECASE)
    
    txt = re.sub(r'(\s|^)\.\s', ' ', txt)
    return txt.strip()

def verify_address(text: str) -> FactualCheck:
    """Uses Nominatim OSM API to verify if physical addresses mentioned actually exist."""
    try:
        # 1. Extraction (Addresses + Context)
        # Use spaCy NER to find GPE/LOC entities instead of unreliable regex
        geo_entities = []
        try:
            from src.preprocessor import _load_spacy
            nlp = _load_spacy()
            doc = nlp(text)
            for ent in doc.ents:
                if ent.label_ in ("GPE", "LOC", "FAC") and len(ent.text.strip()) > 2:
                    geo_entities.append(ent.text.strip())
        except Exception:
            pass  # Fall back to regex below if NER unavailable

        # Also check for explicit street address patterns
        address_patterns = [
            r'\b\d{1,6}\s+[A-Z][a-z\s]+(?:Street|Avenue|St|Ave|Rd|Road|Blvd|Boulevard|Lane|Ln|Drive|Dr|Plaza|Sq|Way|Circle|Cir|Pkwy|Parkway|Court|Ct)\b',
            r'\b[A-Z][a-z\s]+(?:Road|Rd|Way|Lane|Ln|Drive|Dr|Pkwy|Circle|Cir|Street|Avenue|St|Ave)\b'
        ]

        extracted_addresses = []
        for p in address_patterns:
            matches = re.findall(p, text)
            extracted_addresses.extend(matches)

        # If NER found no geographic entities AND regex found no addresses, skip
        if not geo_entities and not extracted_addresses:
            return FactualCheck(check_type="Spatial Validation", result="No geographic entities found to verify.", status="unknown")

        # Context extraction (City, State, Country) — only from NER entities
        valid_context = geo_entities

        # Fallback to regex context if NER found nothing
        if not valid_context:
            geo_context = re.findall(r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\s*,\s*[A-Z]{2,}\b|\b[A-Z][a-z]{3,}\b', text)
            stop_words = {"The", "And", "Today", "News", "Report", "Institute", "Technology", "College", "Francis", "Sardar"}
            valid_context = [g for g in geo_context if g not in stop_words and len(g) > 2]

        if not extracted_addresses:
            if not valid_context:
                 return FactualCheck(check_type="Spatial Validation", result="No specific addresses found to map.", status="unknown")
            query = valid_context[0]
        else:
            query = extracted_addresses[0]
            if valid_context:
                query += f", {valid_context[0]}"

        # 2. OSM Query
        headers = {"User-Agent": "TruthLens_AI_FactChecker/1.0"}
        url = f"https://nominatim.openstreetmap.org/search?q={query}&format=json&limit=1"
        
        response = requests.get(url, headers=headers, timeout=5)
        if response.status_code == 200:
            results = response.json()
            if results:
                display_name = results[0].get("display_name", "")
                has_loc_context = len(valid_context) > 0
                
                # 3. SPATIAL INTEGRITY CHECK
                # Landmark Priority: If we found a named institution, we trust the map name match over House Numbers.
                landmark_keywords = ["Institute", "University", "College", "Hospital", "Temple", "Station", "Airport", "Statue", "National Park"]
                is_landmark = any(lk in query for lk in landmark_keywords) or any(lk in display_name for lk in landmark_keywords)

                if is_landmark:
                    # For landmarks, we just need the city or state to match to avoid false conflicts
                    if valid_context:
                        ctx = valid_context[0].lower()
                        if ctx in display_name.lower():
                             return FactualCheck(
                                check_type="Spatial Validation", 
                                result=f"Verified: Landmark '{query}' confirmed in {valid_context[0]}.", 
                                status="verified"
                            )

                if not valid_context:
                    return FactualCheck(
                        check_type="Spatial Validation", 
                        result=f"Unknown: Entity '{query}' exists, but no specific physical address was mapped.", 
                        status="unknown"
                    )

                if valid_context:
                    text_keywords = set([k.lower() for k in " ".join(valid_context).replace(",", " ").split()])
                    map_keywords = set([k.lower() for k in display_name.replace(",", " ").split()])
                    
                    if not text_keywords.intersection(map_keywords):
                        return FactualCheck(
                            check_type="Spatial Validation", 
                            result=f"Conflict: Result found in {display_name.split(',')[-1]} but claim is for different location.", 
                            status="conflict"
                        )
                
                # House number check (if provided in text)
                house_num = re.search(r'\b\d{1,6}\b', query)
                if house_num and house_num.group() not in display_name:
                     return FactualCheck(
                        check_type="Spatial Validation", 
                        result=f"Conflict: Street found, but house number '{house_num.group()}' is invalid.", 
                        status="conflict"
                    )

                return FactualCheck(
                    check_type="Spatial Validation", 
                    result=f"Verified: Physical location confirmed on global map.", 
                    status="verified"
                )
            else:
                return FactualCheck(
                    check_type="Spatial Validation", 
                    result=f"Conflict: The location '{query}' could not be found on any global maps.", 
                    status="conflict"
                )
        return FactualCheck(check_type="Spatial Validation", result="Geography Index offline. Validation skipped.", status="unknown")
        
    except Exception as e:
        print(f"[RAG] Address error: {e}")
        return FactualCheck(check_type="Spatial Validation", result="Geography Engine error.", status="unknown")

def verify_web_rag(text: str) -> FactualCheck:
    """Enhanced multi-source web verification with credibility scoring and cross-referencing."""
    try:
        queries = _extract_search_queries(text)
        if not queries:
            return FactualCheck(check_type="Web Cross-Reference", result="No cohesive claim found to verify.", status="unknown")

        all_results = []
        seen_urls = set()

        for query in queries:
            try:
                results = _ddgs_search(query, max_results=8)
                for r in results:
                    url = r.get('href', '')
                    if url not in seen_urls:
                        seen_urls.add(url)
                        tier, score = _get_domain_credibility(url)
                        r['_tier'] = tier
                        r['_credibility'] = score
                        all_results.append(r)
            except Exception:
                continue

        if not all_results:
            return FactualCheck(check_type="Web Cross-Reference",
                                result=f"Unconfirmed: No results found across {len(queries)} search queries.",
                                status="unknown")

        all_results.sort(key=lambda r: r['_credibility'], reverse=True)
        combined_web_text = " ".join([r.get('body', '') for r in all_results]).lower()
        main_claim = queries[0].lower()

        # Cross-source consensus: count independent corroborating/conflicting domains
        corroborating_domains = set()
        conflicting_domains = set()
        claim_words = set(main_claim.split()) - _STOPWORDS

        for r in all_results:
            body = r.get('body', '').lower()
            body_words = set(body.split()) - _STOPWORDS
            overlap = claim_words & body_words
            domain = _get_domain(r.get('href', ''))

            if len(claim_words) > 0 and len(overlap) >= 3 and len(overlap) / len(claim_words) > 0.35:
                corroborating_domains.add(domain)

            debunk_kw = ["false", "fake", "hoax", "debunked", "misleading", "fabricated",
                         "not true", "misinformation", "conspiracy theory", "conspiracy",
                         "disproven", "baseless", "unfounded", "pseudoscience", "no evidence",
                         "myth", "rumor", "rumour", "unsubstantiated", "discredited"]
            if any(kw in body for kw in debunk_kw):
                conflicting_domains.add(domain)

        # Numeric fact-check
        stat_patterns = re.findall(r'(?:won|has|have|scored|killed|earned|lost|gained|received)\s+(\d+)', main_claim)
        zero_claim = re.search(r'\b(?:no|zero|0)\s+(?:people|deaths|cases|incidents|victims|casualties)', main_claim)

        if stat_patterns:
            claimed_val = int(stat_patterns[0])
            if claimed_val > 0 and str(claimed_val) in combined_web_text:
                return FactualCheck(check_type="Web Cross-Reference",
                                    result=f"Verified: {len(corroborating_domains)} sources corroborate the count of {claimed_val}.",
                                    status="verified", url=all_results[0].get('href', ''))

        if zero_claim:
            found_counts = re.findall(r'(\d+)\s+(?:people|deaths|cases|incidents|victims|casualties)', combined_web_text)
            for fc in found_counts:
                if int(fc) > 0:
                    return FactualCheck(check_type="Web Cross-Reference",
                                        result=f"Conflict: Multiple sources report {fc} instances, contradicting the zero claim.",
                                        status="conflict", url=all_results[0].get('href', ''))

        # Also check combined text for strong debunking signals
        debunk_combined = ["conspiracy theory", "has been debunked", "no evidence", "false claim",
                           "misinformation", "disproven", "pseudoscience"]
        if any(kw in combined_web_text for kw in debunk_combined):
            # Count how many results contain debunking language
            debunk_count = sum(1 for r in all_results
                               if any(kw in (r.get('body', '') + " " + r.get('title', '')).lower()
                                      for kw in debunk_combined))
            if debunk_count >= 2:
                conflicting_domains.add("__combined_debunk__")

        n_corr = len(corroborating_domains)
        n_confl = len(conflicting_domains)
        best_url = all_results[0].get('href', '')
        best_tier = all_results[0].get('_tier', 'unranked')

        if n_confl >= 2:
            return FactualCheck(check_type="Web Cross-Reference",
                                result=f"Conflict: {n_confl} independent sources flag this claim as false/misleading.",
                                status="conflict", url=best_url)

        if n_corr >= 3:
            tier_note = f" (includes {best_tier} sources)" if best_tier in ("tier1", "tier2") else ""
            return FactualCheck(check_type="Web Cross-Reference",
                                result=f"Verified: {n_corr} independent sources corroborate this claim{tier_note}.",
                                status="verified", url=best_url)
        elif n_corr >= 1:
            if best_tier == "tier1":
                return FactualCheck(check_type="Web Cross-Reference",
                                    result=f"Verified: Corroborated by trusted source ({_get_domain(best_url)}).",
                                    status="verified", url=best_url)
            return FactualCheck(check_type="Web Cross-Reference",
                                result=f"Partially verified: {n_corr} source(s) show overlap but insufficient for full consensus.",
                                status="unknown", url=best_url)

        return FactualCheck(check_type="Web Cross-Reference",
                            result=f"Unconfirmed: Searched {len(all_results)} results across {len(queries)} queries — no strong corroboration.",
                            status="unknown")
    except Exception as e:
        print(f"[RAG] Web search error: {e}")
        return FactualCheck(check_type="Web Cross-Reference", result="Web Search API failure.", status="unknown")

def verify_news_sources(text: str) -> FactualCheck:
    """Search recent news articles for corroboration using DuckDuckGo News."""
    try:
        sentences = [s.strip() for s in re.split(r'[.!?]\s+', text) if len(s.strip()) > 10]
        if not sentences:
            return FactualCheck(check_type="News Corroboration", result="No claim found to verify in news.", status="unknown")

        query = sentences[0][:150]

        time.sleep(1)  # Rate-limit guard
        try:
            news_results = _ddgs_search(query + " news report", max_results=8)
        except Exception:
            news_results = []

        if not news_results:
            return FactualCheck(check_type="News Corroboration",
                                result="No recent news articles found covering this claim.",
                                status="unknown")

        trusted_hits = []
        all_hits = []
        claim_words = set(query.lower().split()) - _STOPWORDS

        for r in news_results:
            body = r.get('body', '').lower()
            title = r.get('title', '').lower()
            combined = body + " " + title
            combined_words = set(combined.split()) - _STOPWORDS
            overlap = claim_words & combined_words

            url = r.get('url', r.get('href', ''))
            domain = _get_domain(url)
            tier, _ = _get_domain_credibility(url)

            if len(claim_words) > 0 and len(overlap) >= 3 and len(overlap) / len(claim_words) > 0.3:
                all_hits.append({"domain": domain, "tier": tier, "url": url})
                if tier in ("tier1", "tier2"):
                    trusted_hits.append({"domain": domain, "tier": tier, "url": url})

        if trusted_hits:
            sources = ", ".join(set(h['domain'] for h in trusted_hits[:3]))
            return FactualCheck(check_type="News Corroboration",
                                result=f"Verified: Covered by {len(trusted_hits)} trusted news source(s): {sources}.",
                                status="verified", url=trusted_hits[0]['url'])
        elif all_hits:
            sources = ", ".join(set(h['domain'] for h in all_hits[:3]))
            return FactualCheck(check_type="News Corroboration",
                                result=f"Partially covered by {len(all_hits)} news source(s): {sources}.",
                                status="unknown", url=all_hits[0]['url'])

        return FactualCheck(check_type="News Corroboration",
                            result="No matching news coverage found for this claim.",
                            status="unknown")
    except Exception as e:
        print(f"[RAG] News search error: {e}")
        return FactualCheck(check_type="News Corroboration", result="News search failed.", status="unknown")


def verify_factcheckers(text: str) -> FactualCheck:
    """Search dedicated fact-checking organizations for verdicts on this claim."""
    try:
        sentences = [s.strip() for s in re.split(r'[.!?]\s+', text) if len(s.strip()) > 10]
        if not sentences:
            return FactualCheck(check_type="Fact-Checker Verdict", result="No claim to verify.", status="unknown")

        query = sentences[0][:100]
        factcheck_sites = ["snopes.com", "politifact.com", "factcheck.org", "fullfact.org", "altnews.in", "boomlive.in"]

        all_fc_results = []
        time.sleep(1)  # Rate-limit guard
        try:
            results = _ddgs_search(query + " fact check", max_results=8)
            for r in results:
                url = r.get('href', '')
                domain = _get_domain(url)
                if any(fc in domain for fc in factcheck_sites):
                    all_fc_results.append(r)
        except Exception:
            pass

        # Fallback: site-specific searches
        if not all_fc_results:
            for site in factcheck_sites[:3]:
                try:
                    results = _ddgs_search(f"site:{site} {query[:80]}", max_results=2)
                    all_fc_results.extend(results)
                except Exception:
                    continue

        if not all_fc_results:
            return FactualCheck(check_type="Fact-Checker Verdict",
                                result="No fact-checker has reviewed this specific claim.",
                                status="unknown")

        for r in all_fc_results:
            body = (r.get('body', '') + " " + r.get('title', '')).lower()
            url = r.get('href', '')
            domain = _get_domain(url)

            false_v = ["false", "pants on fire", "mostly false", "fake", "hoax",
                       "debunked", "misleading", "fabricated", "not true", "unfounded",
                       "conspiracy", "no evidence", "baseless", "disproven", "pseudoscience",
                       "does not cause", "does not", "no,", "is not", "are not",
                       "there is no", "there are no", "did not", "was not", "were not"]
            true_v = ["true", "mostly true", "correct", "verified", "confirmed", "accurate"]
            mixed_v = ["half true", "mixture", "partially true", "unproven", "outdated"]

            if any(v in body for v in false_v):
                return FactualCheck(check_type="Fact-Checker Verdict",
                                    result=f"Conflict: Flagged as FALSE by {domain}.",
                                    status="conflict", url=url)
            elif any(v in body for v in true_v):
                return FactualCheck(check_type="Fact-Checker Verdict",
                                    result=f"Verified: Confirmed as TRUE by {domain}.",
                                    status="verified", url=url)
            elif any(v in body for v in mixed_v):
                return FactualCheck(check_type="Fact-Checker Verdict",
                                    result=f"Mixed verdict from {domain}: claim is partially true or unproven.",
                                    status="unknown", url=url)

        # If a fact-checking site has an article about this topic, it's almost always a debunk
        # (fact-checkers write articles specifically to address false claims)
        fc_domain = _get_domain(all_fc_results[0].get('href', ''))
        is_known_factchecker = any(fc in fc_domain for fc in factcheck_sites)
        if is_known_factchecker:
            return FactualCheck(check_type="Fact-Checker Verdict",
                                result=f"Conflict: Reviewed by {fc_domain} (fact-checkers typically address false claims).",
                                status="conflict", url=all_fc_results[0].get('href', ''))

        return FactualCheck(check_type="Fact-Checker Verdict",
                            result=f"Reviewed by {fc_domain}, but verdict unclear.",
                            status="unknown", url=all_fc_results[0].get('href', ''))
    except Exception as e:
        print(f"[RAG] Fact-checker search error: {e}")
        return FactualCheck(check_type="Fact-Checker Verdict", result="Fact-checker search failed.", status="unknown")


def verify_google_factcheck(text: str) -> FactualCheck:
    """Calls Google Fact Check Tools API (parity with Google Fact Check Explorer).

    Uses the public claims:search endpoint. The endpoint accepts requests
    without an API key for low-volume use; if Google ever gates this, we
    surface "unknown" rather than failing the request.

    Reference:
        https://developers.google.com/fact-check/tools/api
    """
    try:
        # Build a focused query — first sentence (or 120 chars) is usually enough
        sentences = [s.strip() for s in re.split(r'[.!?]\s+', text) if len(s.strip()) > 10]
        query = sentences[0][:120] if sentences else text[:120]
        if not query.strip():
            return FactualCheck(check_type="Google Fact Check",
                                result="No claim text to verify.", status="unknown")

        url = "https://factchecktools.googleapis.com/v1alpha1/claims:search"
        params = {"query": query, "languageCode": "en"}
        if GOOGLE_FACTCHECK_API_KEY:
            params["key"] = GOOGLE_FACTCHECK_API_KEY
        try:
            resp = requests.get(url, params=params, timeout=6)
        except Exception:
            return FactualCheck(check_type="Google Fact Check",
                                result="Google Fact Check API unreachable.", status="unknown")

        if resp.status_code == 429:
            return FactualCheck(check_type="Google Fact Check",
                                result="Google Fact Check rate-limited; try again later.",
                                status="unknown")
        if resp.status_code != 200:
            return FactualCheck(check_type="Google Fact Check",
                                result=f"Google Fact Check returned HTTP {resp.status_code}.",
                                status="unknown")

        data = resp.json() or {}
        claims = data.get("claims") or []
        if not claims:
            return FactualCheck(check_type="Google Fact Check",
                                result="No matching claim in Google's Fact Check index.",
                                status="unknown")

        # Inspect the first matched claim's reviewer verdict
        top = claims[0]
        reviews = top.get("claimReview") or []
        if not reviews:
            return FactualCheck(check_type="Google Fact Check",
                                result=f"Google indexed claim but no rated review found.",
                                status="unknown")

        rv = reviews[0]
        rating = (rv.get("textualRating") or "").strip()
        publisher = ((rv.get("publisher") or {}).get("name")) or "fact-checker"
        review_url = rv.get("url")

        rating_low = rating.lower()
        false_kw = ["false", "pants on fire", "fake", "incorrect", "misleading",
                    "fabricated", "no evidence", "debunked"]
        true_kw = ["true", "correct", "accurate", "verified"]
        mixed_kw = ["mixture", "half true", "partially", "outdated", "unproven"]

        if any(k in rating_low for k in false_kw):
            return FactualCheck(check_type="Google Fact Check",
                                result=f"Conflict: Rated '{rating}' by {publisher}.",
                                status="conflict", url=review_url)
        if any(k in rating_low for k in true_kw):
            return FactualCheck(check_type="Google Fact Check",
                                result=f"Verified: Rated '{rating}' by {publisher}.",
                                status="verified", url=review_url)
        if any(k in rating_low for k in mixed_kw):
            return FactualCheck(check_type="Google Fact Check",
                                result=f"Mixed: Rated '{rating}' by {publisher}.",
                                status="unknown", url=review_url)

        return FactualCheck(check_type="Google Fact Check",
                            result=f"Reviewed by {publisher}: '{rating or 'see source'}'.",
                            status="unknown", url=review_url)
    except Exception as e:
        print(f"[RAG] Google Fact Check error: {e}")
        return FactualCheck(check_type="Google Fact Check",
                            result="Google Fact Check lookup failed.", status="unknown")


# ───────────────────────────────────────────────────────────────────────────
# 8th verifier — LLM Plausibility (Gemini 2.5 Flash)
# ───────────────────────────────────────────────────────────────────────────
# The other 7 verifiers depend on the claim having a fact-checkable surface:
# a date to validate, a person/place on Wikipedia, a Snopes article, a Google
# Fact Check rating. Out-of-domain "scientific breakthrough" hype (miracle
# batteries, instant cures, perpetual motion) routinely returns 7×unknown,
# leaving the stylometric ML model unconstrained.
#
# Gemini fills that gap with world knowledge. Strict label set (EXTRAORDINARY
# / PLAUSIBLE / UNCERTAIN), thinking disabled (this is a label task — we
# don't want Flash to spend output tokens on chain-of-thought before the
# verdict), temperature 0 for determinism. Cached via cache_get/cache_set
# so identical text doesn't re-bill.
_LLM_PLAUSIBILITY_SYSTEM_PROMPT = """You are a scientific plausibility checker for a fake-news detection system.

Reply with EXACTLY two lines, in this exact format. No preamble, no markdown, no extra commentary:

verdict: <LABEL>
reason: <one sentence, 25 words or fewer>

<LABEL> MUST be EXACTLY ONE of these three strings — never invent new labels:

- EXTRAORDINARY — claim violates established science, asserts breakthroughs orders-of-magnitude beyond known precedent without peer-reviewed evidence, or matches the profile of viral fake-tech news (miracle batteries, instant cures, perpetual motion, room-temperature superconductors at trivial cost, "completely replace X within a year"-style commercial timelines that ignore manufacturing reality).
- PLAUSIBLE — claim is consistent with current science, well-attested incremental progress, or routine factual / political / business reporting with no extraordinary technical assertions.
- UNCERTAIN — insufficient information, mixed signals, or topic outside well-established scientific consensus.

If you cannot judge confidently, return UNCERTAIN. Never guess between EXTRAORDINARY and PLAUSIBLE."""


def _parse_llm_verdict(raw: str) -> tuple[str, str]:
    """Pull (verdict, reason) from Gemini's two-line response. Falls back gracefully."""
    if not raw:
        return ("UNCERTAIN", "")
    verdict = "UNCERTAIN"
    reason = ""
    for line in raw.strip().splitlines():
        line = line.strip()
        low = line.lower()
        if low.startswith("verdict:"):
            tok = line.split(":", 1)[1].strip().upper()
            # Be forgiving — strip punctuation/quotes the model may add.
            tok = re.sub(r"[^A-Z]", "", tok)
            if tok in {"EXTRAORDINARY", "PLAUSIBLE", "UNCERTAIN"}:
                verdict = tok
        elif low.startswith("reason:"):
            reason = line.split(":", 1)[1].strip()
    return (verdict, reason)


def verify_llm_plausibility(text: str) -> FactualCheck:
    """Ask Gemini whether the claim is scientifically plausible.

    Maps:
      EXTRAORDINARY → status="conflict"  (Rule 1 fires when factcheck OR LLM flags)
      PLAUSIBLE     → status="verified"
      UNCERTAIN     → status="unknown"
    Always returns a FactualCheck — never raises.
    """
    if not isinstance(text, str) or len(text.strip()) < 30:
        return FactualCheck(check_type="LLM Plausibility",
                            result="Input too short for plausibility check.",
                            status="unknown")

    client = _get_gemini()
    if client is None:
        return FactualCheck(check_type="LLM Plausibility",
                            result="LLM plausibility check unavailable (no GEMINI_API_KEY).",
                            status="unknown")

    # Cache hit → free, fast.
    import hashlib
    cache_key = hashlib.md5(text.strip().encode()).hexdigest()
    cached = cache_get("llm_plausibility", cache_key)
    if cached is not None:
        return FactualCheck(**cached)

    try:
        from google.genai import types as _gtypes
        resp = client.models.generate_content(
            model=LLM_PLAUSIBILITY_MODEL,
            contents=text[:4000],  # paragraph-scale; cap to keep latency tight
            config=_gtypes.GenerateContentConfig(
                system_instruction=_LLM_PLAUSIBILITY_SYSTEM_PROMPT,
                temperature=0.0,
                max_output_tokens=200,
                # Disable thinking — this is a label-extraction task, not reasoning.
                # Default thinking budget would eat the output budget before the verdict prints.
                thinking_config=_gtypes.ThinkingConfig(thinking_budget=0),
            ),
        )
        raw = (resp.text or "").strip()
    except Exception as e:
        print(f"[LLM] Gemini call failed: {e}")
        return FactualCheck(check_type="LLM Plausibility",
                            result="LLM plausibility check failed (API error).",
                            status="unknown")

    verdict, reason = _parse_llm_verdict(raw)
    if verdict == "EXTRAORDINARY":
        check = FactualCheck(
            check_type="LLM Plausibility",
            result=f"Conflict: Claim flagged as scientifically extraordinary. {reason}".strip(),
            status="conflict",
        )
    elif verdict == "PLAUSIBLE":
        check = FactualCheck(
            check_type="LLM Plausibility",
            result=f"Verified: Claim is plausible. {reason}".strip(),
            status="verified",
        )
    else:
        check = FactualCheck(
            check_type="LLM Plausibility",
            result=f"Uncertain: {reason}" if reason else "Uncertain: model returned no clear verdict.",
            status="unknown",
        )

    cache_set("llm_plausibility", cache_key, check.dict())
    return check


_URL_REGEX = re.compile(r'https?://[^\s)>\]"]+', re.IGNORECASE)

def extract_source_credibility(text: str) -> Optional[dict]:
    """If `text` contains a URL, resolve its NewsGuard-style tier + score.

    Returns a SourceCredibility-shaped dict, or None when no URL is present.
    """
    if not isinstance(text, str):
        return None
    match = _URL_REGEX.search(text)
    if not match:
        return None
    found_url = match.group(0)
    tier, score = _get_domain_credibility(found_url)
    domain = _get_domain(found_url)
    if not domain:
        return None
    return {"domain": domain, "tier": tier, "score": float(score)}


def process_document(file: UploadFile, file_bytes: bytes) -> str:
    """Takes RAW uploaded bytes and physically extracts the text inside using PyPDF or Tesseract OCR."""
    filename = file.filename.lower()
    
    # 1. PDF Parsing
    if filename.endswith(".pdf"):
        try:
            reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            text = ""
            for page in reader.pages:
                text += (page.extract_text() or "") + "\n"
            return text.strip()
        except Exception as e:
            print(f"PDF Parse Error: {e}")
            return ""
            
    # 2. Image OCR Parsing
    elif filename.endswith((".png", ".jpg", ".jpeg", ".bmp", ".tiff")):
        if ocr_reader is None:
            # Check if it was just slow to load
            print("OCR Request received but engine not ready. Initializing now...")
            initialize_ocr()
            if ocr_reader is None:
                return "ERROR: OCR_NOT_READY"
                
        try:
            image = Image.open(io.BytesIO(file_bytes)).convert("RGB")
            import numpy as np
            np_img = np.array(image)
            # Execute Neural Text Detection over the image pixels
            results = ocr_reader.readtext(np_img, detail=0)
            text = " ".join(results)
            return text.strip()
        except Exception as e:
            print(f"OCR Error: {e}")
            return ""
    
    # 3. DOCX (Word Document) Parsing
    elif filename.endswith(".docx"):
        try:
            from docx import Document
            doc = Document(io.BytesIO(file_bytes))
            text = ""
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
            return text.strip()
        except Exception as e:
            print(f"DOCX Parse Error: {e}")
            return ""
            
    else:
        # Fallback to absolute raw text string
        try:
            return file_bytes.decode('utf-8')
        except UnicodeDecodeError:
             return ""

@app.get("/")
def read_root():
    return {"status": "ok", "message": "Truth API is running."}

@app.post("/api/analyze", response_model=AnalyzeResponse)
async def analyze_article(
    text: Optional[str] = Form(None),
    file: Optional[UploadFile] = File(None)
):
    if not text and not file:
        raise HTTPException(status_code=400, detail="Must provide either text or a file.")
    
    start_time = time.time()
    
    parsed_text = ""
    if file:
        file_bytes = await file.read()
        extracted = process_document(file, file_bytes)
        
        if extracted == "ERROR: OCR_NOT_READY":
             raise HTTPException(status_code=503, detail="The AI OCR engine is currently initializing (downloading neural weights). Please wait 30 seconds and try again.")
             
        if len(extracted.strip()) > 0:
            parsed_text = extracted
        else:
            raise HTTPException(status_code=400, detail="The OCR engine could not detect any readable text in this document. Please ensure the image is clear and contains text.")
    
    if text:
        parsed_text += " " + text.strip()
        
    # ---------------------------
    # TRUTHLENS ML ENSEMBLE PREDICTION
    # ---------------------------
    is_fake = False
    prediction = "Real News"
    confidence = 0.50
    
    vocab_coverage = 0.0  # Track how much of the input the model actually understands

    if feature_engine and ensemble_model and len(parsed_text.strip()) > 0:
        try:
            # First, clean, mask entities, and lemmatize using TruthLens built-in NLP.
            # This perfectly aligns the user's text with the TF-IDF Vocabulary matrices.
            df = pd.DataFrame([{"text": parsed_text}])
            processed_df = preprocess_dataframe(df)

            if len(processed_df) == 0:
                raise ValueError("Text was entirely stripped during NLP processing.")

            p_text = processed_df["processed_text"].tolist()
            c_text = processed_df["cleaned_text"].tolist()
            r_text = processed_df["text"].tolist()

            # --- VOCABULARY COVERAGE CHECK ---
            # Measure how many input words are in the TF-IDF vocabulary.
            # If coverage is very low, the model has no content signal and
            # the prediction would be driven entirely by stylometric features,
            # which is unreliable for non-news text.
            tfidf_vec = feature_engine.tfidf.transform(p_text).toarray()[0]
            input_word_count = max(len(p_text[0].split()), 1)
            matched_terms = int(np.count_nonzero(tfidf_vec))
            vocab_coverage = matched_terms / input_word_count  # ratio of input words recognized
            print(f"[ML] Vocab coverage: {matched_terms}/{input_word_count} words ({vocab_coverage:.0%})")

            # Convert text into the TruthLens feature geometry
            X_features = feature_engine.transform(p_text, c_text, r_text)

            # Feed feature matrix into the Logistic Regression Meta-Learner Stack
            probs = ensemble_model.predict_proba(X_features)[0]

            # probs[0] = Real (Label 0), probs[1] = Fake (Label 1)
            if probs[1] > probs[0]:
                is_fake = True
                prediction = "Fake News"
                confidence = round(float(probs[1]), 2)
            else:
                is_fake = False
                prediction = "Real News"
                confidence = round(float(probs[0]), 2)

            # --- CONFIDENCE ADJUSTMENT (vocabulary-coverage abstention gate) ---
            # The model needs sufficient content signal (TF-IDF matches) to make
            # a reliable prediction. With too few recognized words, the prediction
            # is driven by stylometric features alone, which is unreliable.
            #
            # Thresholds come from TruthLens config (single source of truth so
            # the analysis script in scripts/selective_risk.py and this gate
            # can never drift):
            #   _MIN_MATCHED_TERMS         (default 4)
            #   _VOCAB_COVERAGE_THRESHOLD  (default 0.30)
            insufficient_data = (
                matched_terms < _MIN_MATCHED_TERMS
                or vocab_coverage < _VOCAB_COVERAGE_THRESHOLD
            )
            if insufficient_data:
                # Scale confidence by how much content the model actually understood
                signal_strength = min(
                    matched_terms / max(_MIN_MATCHED_TERMS, 1),
                    vocab_coverage / max(_VOCAB_COVERAGE_THRESHOLD, 1e-6),
                )
                adjusted = confidence * signal_strength
                print(f"[ML] Low signal: {matched_terms} terms, {vocab_coverage:.0%} coverage. "
                      f"Confidence {confidence} -> {adjusted:.2f}")
                confidence = round(adjusted, 2)
                if confidence < 0.55:
                    prediction = "Uncertain"
                    is_fake = False

        except Exception as e:
            print(f"ML Transform Error: {e}")
            # Fall through to heuristic — never return "Error" as prediction
            prediction = "Error"

    # Heuristic fallback: used when ML models not loaded OR when ML fails on short text
    if prediction == "Error" or (not feature_engine or not ensemble_model):
        sensational_words = ['!!!', 'shocking', 'breaking', 'coverup', 'hoax',
                             'conspiracy', 'they dont want you to know', 'wake up',
                             'exposed', 'banned', 'censored', 'share before deleted']
        exclamation_heavy = parsed_text.count('!') >= 3
        all_caps_ratio = sum(1 for c in parsed_text if c.isupper()) / max(len(parsed_text), 1)
        has_sensational = any(w in parsed_text.lower() for w in sensational_words)

        if has_sensational or exclamation_heavy or all_caps_ratio > 0.4:
            is_fake = True
            prediction = "Fake News"
            confidence = 0.65  # Low confidence since heuristic-only
        else:
            is_fake = False
            prediction = "Real News"
            confidence = 0.50  # Uncertain — let RAG decide
        print(f"HEURISTIC FALLBACK: {prediction} (confidence={confidence})")
        
    # Stash the raw ML verdict before any RAG override fires — the
    # ConflictReport panel reports it side-by-side with the aggregated
    # RAG verdict so users can see when the two disagreed.
    ml_verdict_pre_rag = prediction
    ml_confidence_pre_rag = float(confidence)
    bias_gate_active = bool(
        feature_engine and ensemble_model
        and vocab_coverage < _VOCAB_COVERAGE_THRESHOLD
    )

    # ---------------------------
    # RAG FACTUAL ANALYSIS
    # ---------------------------
    factual_analysis = []
    conflict_detected = False

    # Clean OCR text of UI noise (like "See outside" or "Google Reviews")
    # to prevent RAG from searching for buttons instead of subjects.
    clean_rag_text = scrub_noise(parsed_text)
    
    temporal_result = verify_timeline(clean_rag_text)
    factual_analysis.append(temporal_result)
    if temporal_result.status == "conflict":
        conflict_detected = True

    web_result = verify_web_rag(clean_rag_text)
    factual_analysis.append(web_result)
    if web_result.status == "conflict":
        conflict_detected = True

    wiki_result = verify_wikipedia(clean_rag_text)
    factual_analysis.append(wiki_result)
    if wiki_result.status == "conflict":
        conflict_detected = True

    news_result = verify_news_sources(clean_rag_text)
    factual_analysis.append(news_result)
    if news_result.status == "conflict":
        conflict_detected = True

    factcheck_result = verify_factcheckers(clean_rag_text)
    factual_analysis.append(factcheck_result)
    if factcheck_result.status == "conflict":
        conflict_detected = True

    address_result = verify_address(clean_rag_text)
    factual_analysis.append(address_result)
    if address_result.status == "conflict":
        verified_other = [c for c in factual_analysis[:-1] if c.status == "verified"]
        if verified_other:
            address_result.status = "unknown"
            address_result.result = address_result.result.replace("Conflict:", "Unreliable (Subject Verified elsewhere):")
        else:
            conflict_detected = True

    # 7th verifier — Google Fact Check Tools API (parity with Google Fact Check Explorer).
    # Treated alongside the other fact-checkers in the override rules below.
    google_fc_result = verify_google_factcheck(clean_rag_text)
    factual_analysis.append(google_fc_result)
    if google_fc_result.status == "conflict":
        conflict_detected = True

    # 8th verifier — LLM Plausibility (Gemini 2.5 Flash). World-knowledge
    # check that catches OOD claims the other 7 verifiers can't, e.g. miracle
    # batteries / instant cures / perpetual motion. EXTRAORDINARY → conflict
    # → Rule 1 fires alongside the other fact-checkers.
    llm_result = verify_llm_plausibility(clean_rag_text)
    factual_analysis.append(llm_result)
    if llm_result.status == "conflict":
        conflict_detected = True

    # -- RAG EVIDENCE INTEGRATION (7-SOURCE CONSENSUS) --
    conflict_checks = [c for c in factual_analysis if c.status == "conflict"]
    verified_checks = [c for c in factual_analysis if c.status == "verified"]
    num_conflicts = len(conflict_checks)
    num_verified = len(verified_checks)

    # Fact-checker verdicts carry extra weight. We treat the in-house
    # search-based fact-checker (`factcheck_result`), Google Fact Check Tools
    # (`google_fc_result`), and the LLM plausibility check (`llm_result`) as
    # a coalition — any of them flagging FALSE / EXTRAORDINARY triggers Rule 1;
    # any confirming TRUE / PLAUSIBLE triggers Rule 5.
    fc_conflict = (factcheck_result.status == "conflict"
                   or google_fc_result.status == "conflict"
                   or llm_result.status == "conflict")
    fc_verified = (factcheck_result.status == "verified"
                   or google_fc_result.status == "verified"
                   or llm_result.status == "verified")
    triggered_rule = None  # populated by whichever override rule fires; used by ConflictReport

    # Rule 1: Fact-checker says FALSE → strong override
    if fc_conflict:
        prediction = "Fake News"
        confidence = max(confidence, 0.92)
        is_fake = True
        conflict_detected = True
        triggered_rule = "Rule 1: A fact-checker flagged this claim as false."
        print(f"RAG OVERRIDE: Fact-checker flagged as FALSE. Result set to FAKE (conf={confidence}).")
    # Rule 2: Multiple conflicts across different sources
    elif num_conflicts >= 2:
        prediction = "Fake News"
        confidence = max(confidence, 0.88)
        is_fake = True
        conflict_detected = True
        triggered_rule = f"Rule 2: {num_conflicts} sources reported conflicts."
        print(f"RAG OVERRIDE: {num_conflicts} conflicts from: {[c.check_type for c in conflict_checks]}.")
    # Rule 3: Single conflict — reduce confidence, flag it
    elif num_conflicts == 1:
        conflict_detected = True
        triggered_rule = f"Rule 3: 1 source ({conflict_checks[0].check_type}) reported a conflict."
        if prediction == "Real News":
            confidence = min(confidence, 0.70)
            print(f"RAG WARNING: 1 conflict found ({conflict_checks[0].check_type}). Confidence capped at {confidence}.")
        else:
            confidence = max(confidence, 0.90)
            print(f"RAG CONFIRM: 1 conflict agrees with ML Fake prediction.")
    # Rule 4: Strong multi-source verification
    elif num_verified >= 3:
        if prediction == "Fake News" and confidence < 0.85:
            prediction = "Real News"
            confidence = max(confidence, 0.85)
            is_fake = False
            triggered_rule = f"Rule 4: {num_verified} sources verified — overrides low-confidence Fake."
            print(f"RAG OVERRIDE: {num_verified} verifications override low-confidence Fake.")
        elif prediction == "Real News":
            confidence = max(confidence, 0.90)
            triggered_rule = f"Rule 4: {num_verified} sources verified — confidence boosted."
            print(f"RAG CONFIRM: {num_verified} sources verify. Confidence boosted to {confidence}.")
    elif num_verified >= 2:
        if prediction == "Fake News" and confidence < 0.80:
            prediction = "Real News"
            confidence = max(confidence, 0.80)
            is_fake = False
            triggered_rule = f"Rule 4: {num_verified} sources verified — overrides low-confidence Fake."
            print(f"RAG OVERRIDE: {num_verified} verifications override low-confidence Fake.")
        elif prediction == "Real News":
            confidence = max(confidence, 0.85)
            triggered_rule = f"Rule 4: {num_verified} sources support Real prediction."
            print(f"RAG CONFIRM: {num_verified} verifications support Real prediction.")
    # Rule 7: No external evidence at all (zero conflicts AND zero verifications).
    # The model alone is making the call — cap confidence so we never claim
    # 100% on an input no verifier could corroborate.
    elif num_conflicts == 0 and num_verified == 0 and not fc_conflict and not fc_verified:
        if prediction in ("Real News", "Fake News") and confidence > 0.65:
            confidence = 0.65
            triggered_rule = "Rule 7: No external evidence; ML verdict cannot be confirmed."
            print(f"RAG WARNING: No verifier found supporting evidence. Confidence capped at {confidence}.")
    # Rule 5: Fact-checker confirmed TRUE → boost
    if fc_verified and prediction == "Real News":
        confidence = max(confidence, 0.93)
        # Don't overwrite a more specific Rule 1-4 trigger — only annotate if empty.
        if triggered_rule is None:
            triggered_rule = "Rule 5: A fact-checker confirmed this claim as true."
        print(f"RAG BOOST: Fact-checker confirmed TRUE. Confidence={confidence}.")
    # Rule 6: ML failed, use RAG consensus
    if prediction == "Error":
        if num_conflicts > 0:
            prediction = "Fake News"
            confidence = 0.75
            is_fake = True
            triggered_rule = "Rule 6: ML inference failed; RAG consensus → Fake."
        elif num_verified > 0:
            prediction = "Real News"
            confidence = 0.65 + (0.05 * min(num_verified, 4))
            is_fake = False
            triggered_rule = "Rule 6: ML inference failed; RAG consensus → Real."
        else:
            prediction = "Uncertain"
            confidence = 0.50
            is_fake = False
            triggered_rule = "Rule 6: ML inference failed and RAG was inconclusive."
        print(f"RAG FALLBACK: ML failed, RAG decided: {prediction} (conf={confidence})")

    # ---------------------------
    # CONTENT + STYLE EXPLAINABILITY
    # ---------------------------
    # SVD backprojection covers TF-IDF portion (150/265 dims); style indicators appended separately
    # Instead of raw TF-IDF weights (which don't reflect what the model sees),
    # we backproject through the SVD to find which words contribute most to the
    # 150-dimensional features the model actually uses for classification.
    #
    # Method: word_importance = tfidf_vector @ svd.components_.T @ svd.components_
    # This gives each word's contribution through the SVD bottleneck.
    explainability = []
    if feature_engine and len(parsed_text.strip()) > 0:
        try:
            processed_str = p_text[0] if 'p_text' in locals() else parsed_text
            tfidf_vec = feature_engine.tfidf.transform([processed_str]).toarray()[0]
            feature_names = feature_engine.tfidf.get_feature_names_out()

            # Backproject: how much of each word survives through SVD
            # svd.components_ shape: (150, n_features) — each row is a component
            svd_projected = tfidf_vec @ feature_engine.svd.components_.T  # (150,)
            backprojected = svd_projected @ feature_engine.svd.components_  # (n_features,)
            # word_importance = how much each original word contributes to the SVD space
            word_importance = np.abs(backprojected)

            top_indices = word_importance.argsort()[-5:][::-1]

            for idx in top_indices:
                word = feature_names[idx]
                weight = word_importance[idx]
                if weight > 0.001:
                    c = "rgba(244, 63, 94, 0.8)" if is_fake else "rgba(16, 185, 129, 0.8)"
                    w = -round(float(weight), 2) if is_fake else round(float(weight), 2)
                    explainability.append(LimeFeature(word=word, weight=w, color=c))
                if len(explainability) >= 3:
                    break
        except Exception as e:
            print(f"Explainability error: {e}")

    # Add stylometric indicators if we have room (cap total at 5)
    if len(explainability) < 5 and len(parsed_text.strip()) > 0:
        exc_count = parsed_text.count('!')
        cap_ratio = sum(1 for ch in parsed_text if ch.isupper()) / max(len(parsed_text), 1)
        style_color = "rgba(244, 63, 94, 0.8)" if is_fake else "rgba(16, 185, 129, 0.8)"
        if exc_count >= 3 and len(explainability) < 5:
            w = -0.2 if is_fake else 0.1
            explainability.append(LimeFeature(word="[style] excessive punctuation (!)", weight=w, color=style_color))
        if cap_ratio > 0.25 and len(explainability) < 5:
            w = -0.2 if is_fake else 0.1
            explainability.append(LimeFeature(word="[style] heavy capitalization", weight=w, color=style_color))
        if vocab_coverage < _VOCAB_COVERAGE_THRESHOLD and len(explainability) < 5:
            w = -0.15 if is_fake else 0.15
            explainability.append(LimeFeature(word="[style] out-of-domain vocabulary", weight=w, color=style_color))

    # Fallback: if SVD backprojection failed, show stylometric indicators
    if len(explainability) == 0:
        style_indicators = []
        exc_count = parsed_text.count('!')
        cap_ratio = sum(1 for c in parsed_text if c.isupper()) / max(len(parsed_text), 1)
        if exc_count >= 3:
            style_indicators.append(("excessive punctuation", -0.3 if is_fake else 0.1))
        if cap_ratio > 0.3:
            style_indicators.append(("heavy capitalization", -0.3 if is_fake else 0.1))
        if not style_indicators:
            words = parsed_text.split()[:3]
            for w in words:
                if len(w) > 2:
                    style_indicators.append((w.lower(), -0.1 if is_fake else 0.1))
                    break
        c = "rgba(244, 63, 94, 0.8)" if is_fake else "rgba(16, 185, 129, 0.8)"
        for word, weight in style_indicators[:3]:
            explainability.append(LimeFeature(word=word, weight=weight, color=c))
            
    # ---------------------------
    # PARITY SURFACES — claim type, source credibility, sentence scoring
    # ---------------------------
    claim_type_value = None
    if classify_claim_input is not None:
        try:
            ct = classify_claim_input(parsed_text)
            claim_type_value = ct.get("type") if isinstance(ct, dict) else None
        except Exception as e:
            print(f"[claim_type] failed: {e}")

    cred_dict = extract_source_credibility(parsed_text)
    source_credibility_value = SourceCredibility(**cred_dict) if cred_dict else None

    sentence_score_objs: List[SentenceScore] = []
    if (
        score_sentences is not None
        and feature_engine is not None
        and ensemble_model is not None
        and len(parsed_text.strip()) >= 40
    ):
        try:
            from src import preprocessor as _preprocessor_mod
            raw_scores = score_sentences(
                parsed_text, feature_engine, ensemble_model,
                preprocessor_module=_preprocessor_mod,
            )
            for r in raw_scores:
                sentence_score_objs.append(SentenceScore(**r))
        except Exception as e:
            print(f"[sentence_scores] failed: {e}")

    # ---------------------------
    # CONFLICT REPORT — structured ML-vs-RAG breakdown for the UI panel
    # ---------------------------
    rag_verdict_label = "Real News"
    if num_conflicts >= 2 or fc_conflict:
        rag_verdict_label = "Fake News"
    elif num_conflicts == 1:
        rag_verdict_label = "Disputed"
    elif num_verified >= 2:
        rag_verdict_label = "Real News"
    elif num_conflicts == 0 and num_verified == 0:
        rag_verdict_label = "Unknown"
    rag_total_signal = max(num_conflicts + num_verified, 1)
    rag_confidence_value = round(num_verified / rag_total_signal, 2) if rag_verdict_label == "Real News" \
        else round(num_conflicts / rag_total_signal, 2) if rag_verdict_label == "Fake News" \
        else 0.5
    ml_disagreement = (
        ml_verdict_pre_rag in ("Real News", "Fake News")
        and rag_verdict_label in ("Real News", "Fake News")
        and ml_verdict_pre_rag != rag_verdict_label
    )
    if triggered_rule and ("override" in triggered_rule.lower() or "Rule 1" in triggered_rule or "Rule 2" in triggered_rule):
        winning_signal = "rag"
    elif bias_gate_active:
        winning_signal = "consensus"
    elif ml_disagreement:
        winning_signal = "rag" if triggered_rule else "ml"
    else:
        winning_signal = "ml"
    flagging_verifier_names = [c.check_type for c in conflict_checks]

    conflict_report_value = ConflictReport(
        ml_verdict=ml_verdict_pre_rag,
        ml_confidence=round(ml_confidence_pre_rag, 2),
        rag_verdict=rag_verdict_label,
        rag_confidence=rag_confidence_value,
        disagreement=ml_disagreement,
        winning_signal=winning_signal,
        triggered_rule=triggered_rule,
        flagging_verifiers=flagging_verifier_names,
        top_lime_tokens=explainability[:5],
        vocab_coverage=round(vocab_coverage, 2),
        bias_gate_active=bias_gate_active,
    )

    processing_time = time.time() - start_time

    return AnalyzeResponse(
        prediction=prediction,
        confidence=confidence,
        processing_time=round(processing_time, 2),
        explainability=explainability,
        analyzed_text=parsed_text,
        factual_analysis=factual_analysis,
        conflict_detected=conflict_detected,
        vocab_coverage=round(vocab_coverage, 2),
        sentence_scores=sentence_score_objs,
        claim_type=claim_type_value,
        source_credibility=source_credibility_value,
        conflict_report=conflict_report_value,
    )

# Serve Frontend (Production Mode)
# This assumes the frontend has been built into Truth/frontend/dist
frontend_path = os.path.join(os.path.dirname(backend_dir), "frontend", "dist")

if os.path.exists(frontend_path):
    app.mount("/", StaticFiles(directory=frontend_path, html=True), name="static")

    @app.exception_handler(404)
    async def not_found_exception_handler(request, exc):
        # Fallback to index.html for Single Page App routing
        return FileResponse(os.path.join(frontend_path, "index.html"))
else:
    print(f"Frontend dist not found at {frontend_path}. API Only mode.")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
